"""
문서 처리 모듈: DOCX, PDF 파일을 텍스트로 추출하고 청크 단위로 분리
"""

import os
from typing import List, Dict, Any, Union
import docx
import fitz  # PyMuPDF
import re
import nltk
from nltk.tokenize import sent_tokenize

# nltk 데이터 필요시 다운로드
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt', quiet=True)

def extract_text(file_path: str) -> str:
    """
    DOCX 또는 PDF 파일에서 텍스트를 추출
    
    Args:
        file_path: 처리할 파일 경로
        
    Returns:
        추출된 전체 텍스트
    """
    file_ext = os.path.splitext(file_path)[1].lower()
    
    if file_ext == '.docx':
        return _extract_from_docx(file_path)
    elif file_ext == '.pdf':
        return _extract_from_pdf(file_path)
    else:
        raise ValueError(f"지원하지 않는 파일 형식입니다: {file_ext}")

def _extract_from_docx(file_path: str) -> str:
    """DOCX 파일에서 텍스트를 추출"""
    try:
        doc = docx.Document(file_path)
        full_text = []
        
        # 디버깅 정보
        print(f"DOCX 파일 처리: {file_path}")
        print(f"단락 수: {len(doc.paragraphs)}")
        
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        # 표(tables)에서 텍스트 추출
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        
        return '\n'.join(full_text)
    except Exception as e:
        print(f"DOCX 파일 처리 중 오류: {str(e)}")
        # 파일 존재 여부 확인
        import os
        if os.path.exists(file_path):
            file_size = os.path.getsize(file_path)
            print(f"파일은 존재함 (크기: {file_size} 바이트)")
        else:
            print(f"파일이 존재하지 않음: {file_path}")
        raise

def _extract_from_pdf(file_path: str) -> str:
    """PDF 파일에서 텍스트를 추출"""
    try:
        doc = fitz.open(file_path)
        full_text = []
        
        # 디버깅 정보
        print(f"PDF 파일 처리: {file_path}")
        print(f"페이지 수: {len(doc)}")
        
        for page in doc:
            full_text.append(page.get_text())
        
        return '\n'.join(full_text)
    except Exception as e:
        print(f"PDF 파일 처리 중 오류: {str(e)}")
        # 파일 존재 여부 확인
        import os
        if os.path.exists(file_path):
            file_size = os.path.getsize(file_path)
            print(f"파일은 존재함 (크기: {file_size} 바이트)")
        else:
            print(f"파일이 존재하지 않음: {file_path}")
        raise

def split_text(text: str, chunk_size: int = 1000, chunk_overlap: int = 200) -> List[str]:
    """
    텍스트를 청크로 분할
    
    Args:
        text: 분할할 텍스트
        chunk_size: 각 청크의 최대 크기
        chunk_overlap: 청크 간 겹치는 문자 수
        
    Returns:
        분할된 텍스트 청크 리스트
    """
    # 텍스트가 비어있는 경우
    if not text:
        return []
    
    # 단락 단위로 분할
    paragraphs = text.split('\n')
    paragraphs = [p for p in paragraphs if p.strip()]
    
    chunks = []
    current_chunk = []
    current_size = 0
    
    for para in paragraphs:
        # 단락이 청크 사이즈보다 크면 추가 분할
        if len(para) > chunk_size:
            # 현재 청크를 저장
            if current_chunk:
                chunks.append('\n'.join(current_chunk))
                # 겹치는 부분 유지
                overlap_paras = current_chunk[-1:] if current_chunk else []
                current_chunk = overlap_paras
                current_size = sum(len(p) for p in current_chunk)
            
            # 큰 단락 분할
            words = para.split()
            temp_chunk = []
            temp_size = 0
            
            for word in words:
                if temp_size + len(word) + 1 > chunk_size:
                    chunks.append(' '.join(temp_chunk))
                    # 겹치는 부분 유지
                    overlap_point = max(0, len(temp_chunk) - int(chunk_overlap / 5))
                    temp_chunk = temp_chunk[overlap_point:]
                    temp_size = sum(len(w) + 1 for w in temp_chunk)
                
                temp_chunk.append(word)
                temp_size += len(word) + 1
            
            if temp_chunk:
                current_chunk.append(' '.join(temp_chunk))
                current_size += temp_size
        
        # 일반적인 경우: 단락 추가
        elif current_size + len(para) > chunk_size:
            chunks.append('\n'.join(current_chunk))
            # 겹치는 부분 유지
            overlap_paras = current_chunk[-1:] if current_chunk else []
            current_chunk = overlap_paras
            current_size = sum(len(p) for p in current_chunk)
            current_chunk.append(para)
            current_size += len(para)
        else:
            current_chunk.append(para)
            current_size += len(para)
    
    # 마지막 청크 추가
    if current_chunk:
        chunks.append('\n'.join(current_chunk))
    
    return chunks

def process_document(file_path: str, chunk_size: int = 1000, chunk_overlap: int = 200) -> List[Dict[str, Any]]:
    """
    문서를 처리하여 청크 단위로 분리
    
    Args:
        file_path: 처리할 파일 경로
        chunk_size: 각 청크의 최대 크기
        chunk_overlap: 청크 간 겹치는 문자 수
        
    Returns:
        청크 리스트 (메타데이터 포함)
    """
    # 텍스트 추출
    text = extract_text(file_path)
    
    # 텍스트 분할
    chunks = split_text(text, chunk_size, chunk_overlap)
    
    # 메타데이터 추가 (파일명, 페이지 번호 등)
    file_name = os.path.basename(file_path)
    processed_chunks = []
    
    for i, chunk_text in enumerate(chunks):
        processed_chunks.append({
            'text': chunk_text,
            'metadata': {
                'file_name': file_name,
                'chunk_id': i,
                'source': file_path
            }
        })
    
    return processed_chunks

# 예시 기획서 데이터 추가 (테스트 용도)
def generate_sample_game_design_doc() -> List[Dict[str, Any]]:
    """
    게임 스킬 시스템 및 아이템 장착 관련 예시 데이터 생성
    
    Returns:
        게임 기획서 청크 목록
    """
    sample_texts = [
        # 스킬 시스템의 아이템 장착 기능 기획서
        """
        ## 스킬 시스템 - 아이템 장착 기능
        
        ### 개요
        스킬 시스템 내 아이템 장착 기능은 캐릭터의 능력치와 스킬 성능에 직접적인 영향을 주는 핵심 메커니즘입니다.
        플레이어는 다양한 장비를 장착하여 캐릭터의 전투력을 향상시키고 새로운 스킬 효과를 활성화할 수 있습니다.
        
        ### 장비 슬롯 구성
        - 무기: 주무기, 보조무기
        - 방어구: 헬멧, 갑옷, 장갑, 신발
        - 장신구: 목걸이, 반지(2개), 귀걸이
        
        ### 장비 장착 효과
        1. 스탯 증가
           - 공격력, 방어력, 체력 등 기본 능력치 증가
           - 특수 능력치(치명타, 회피율 등) 증가
        
        2. 스킬 강화 효과
           - 특정 스킬 데미지 증가 (예: 화염 검 장착 시 화염 스킬 데미지 +20%)
           - 스킬 쿨타임 감소 (예: 민첩의 장갑 장착 시 회피 스킬 쿨타임 -15%)
           - 스킬 지속시간 증가 (예: 집중의 투구 장착 시 버프 지속시간 +5초)
        
        3. 세트 효과
           - 동일 세트 아이템 2개 이상 장착 시 추가 효과 발동
           - 세트 단계별 효과 (2세트, 4세트, 6세트)
           - 세트 완성 시 특수 스킬 해금
        
        ### 장비 장착 제한 조건
        - 캐릭터 레벨 제한 (최소 요구 레벨)
        - 클래스 제한 (특정 직업만 장착 가능)
        - 스탯 요구사항 (최소 요구 스탯)
        
        ### 장비 장착 UI 및 인터랙션
        - 인벤토리에서 장비 슬롯으로 드래그 앤 드롭
        - 장비 아이콘 더블 클릭으로 자동 장착
        - 장착 가능 아이템 하이라이트 표시
        - 장착 불가 아이템은 빨간색으로 표시
        - 장착 효과 비교 툴팁 표시
        
        ### 장비 장착 시각적 피드백
        - 장비 장착 시 캐릭터 외형 변경
        - 특수 효과 아이템 장착 시 시각적 이펙트 추가
        - 세트 효과 활성화 시 특수 오라 효과
        
        ### 장비 변경 및 해제
        - 장비 슬롯 클릭으로 장비 해제
        - 다른 장비로 교체 시 자동 해제 후 장착
        - 장비 변경 시 스킬 UI에 변경 효과 표시
        """,
        
        # 아이템 장착에 따른 스킬 효과 상세 기획서
        """
        ## 아이템 장착에 따른 스킬 효과 상세
        
        ### 무기 타입별 스킬 영향
        1. 검류
           - 근접 공격 스킬 데미지 +15%
           - 연속 공격 스킬 추가 타격 확률 +10%
        
        2. 도끼류
           - 단일 타격 스킬 치명타 확률 +20%
           - 방어력 관통 효과 +10%
        
        3. 창류
           - 공격 범위 +15%
           - 관통 공격 스킬 추가 대상 +1
        
        4. 지팡이류
           - 마법 스킬 데미지 +20%
           - 마나 소모량 -10%
        
        5. 활류
           - 원거리 스킬 명중률 +15%
           - 치명타 데미지 +25%
        
        ### 방어구 세트별 스킬 영향
        1. 전사의 의지(세트)
           - 2세트: 물리 데미지 감소 +10%
           - 4세트: 도발 스킬 지속시간 +3초
           - 6세트: 새로운 방어 스킬 '철벽 방어' 해금
        
        2. 마법사의 통찰(세트)
           - 2세트: 마법 데미지 +15%
           - 4세트: 마나 회복 속도 +20%
           - 6세트: 새로운 공격 스킬 '아케인 폭발' 해금
        
        3. 암살자의 그림자(세트)
           - 2세트: 회피율 +10%
           - 4세트: 은신 스킬 지속시간 +4초
           - 6세트: 새로운 기동 스킬 '그림자 도약' 해금
        
        ### 특수 장비 효과
        1. 속성 부여 장비
           - 무기에 속성 부여 (화염, 빙결, 번개 등)
           - 해당 속성 스킬 사용 시 추가 효과 발동
        
        2. 쿨타임 감소 장비
           - 특정 스킬 카테고리의 쿨타임 감소
           - 전체 스킬 재사용 대기시간 감소
        
        3. 자원 회복 장비
           - 스킬 사용 시 HP/MP 회복
           - 특정 조건 달성 시 자원 즉시 회복
        
        ### 장비 강화 시스템과 스킬 연동
        1. 장비 강화 레벨에 따른 스킬 보너스
           - +1 ~ +5: 기본 스킬 데미지 소폭 증가
           - +6 ~ +10: 스킬 부가 효과 강화
           - +11 ~ +15: 스킬 특수 효과 추가
        
        2. 장비 레벨 업그레이드
           - 특정 재료 수집 시 장비 레벨 상승
           - 레벨 상승에 따른 새로운 스킬 효과 해금
        
        3. 룬 시스템
           - 장비에 룬 장착으로 스킬 커스터마이징
           - 다양한 조합으로 독특한 스킬 효과 생성
        """,
        
        # 아이템 장착 예외 상황 및 오류 처리 기획서
        """
        ## 아이템 장착 예외 상황 및 오류 처리
        
        ### 장착 제한 조건 처리
        1. 레벨 제한
           - 상황: 캐릭터 레벨이 아이템 요구 레벨 미만인 경우
           - 처리: 장착 버튼 비활성화, 요구 레벨 툴팁 표시
           - 메시지: "이 아이템을 장착하려면 레벨 {요구레벨}이 필요합니다"
        
        2. 클래스 제한
           - 상황: 다른 클래스 전용 아이템을 장착하려는 경우
           - 처리: 장착 버튼 비활성화, 요구 클래스 툴팁 표시
           - 메시지: "이 아이템은 {클래스명} 전용입니다"
        
        3. 스탯 제한
           - 상황: 요구 스탯이 부족한 경우
           - 처리: 장착 버튼 비활성화, 부족한 스탯 표시
           - 메시지: "{스탯명} {필요수치} 이상이 필요합니다 (현재: {현재수치})"
        
        ### 장비 슬롯 충돌 처리
        1. 이미 장착된 슬롯
           - 상황: 이미 다른 아이템이 장착된 슬롯에 장착 시도
           - 처리: 확인 다이얼로그 표시
           - 메시지: "이미 장착된 아이템을 교체하시겠습니까?"
        
        2. 양손/한손 무기 충돌
           - 상황: 양손 무기 장착 시 보조 무기/방패 슬롯 충돌
           - 처리: 자동으로 보조 무기/방패 해제 후 확인 요청
           - 메시지: "양손 무기 장착 시 보조 무기가 해제됩니다. 계속하시겠습니까?"
        
        3. 고유 장비 중복 장착
           - 상황: '고유 장착' 제한이 있는 아이템 중복 장착 시도
           - 처리: 장착 불가, 이미 장착된 고유 아이템 하이라이트
           - 메시지: "고유 아이템은 한 번만 장착할 수 있습니다"
        
        ### 특수 상황 처리
        1. 내구도 소진 아이템
           - 상황: 내구도가 0인 파괴된 장비 장착 시도
           - 처리: 장착 불가, 수리 옵션 제시
           - 메시지: "파괴된 아이템은 수리 후 장착할 수 있습니다"
        
        2. 임시 효과 아이템
           - 상황: 시간 제한이 있는 특수 아이템 장착
           - 처리: 남은 시간 툴팁 표시, 시간 경과 시 자동 해제
           - 메시지: "이 아이템은 {남은시간} 동안 사용할 수 있습니다"
        
        3. 이벤트 아이템
           - 상황: 특정 이벤트 기간에만 사용 가능한 아이템
           - 처리: 이벤트 기간 외에는 장착 불가
           - 메시지: "이 아이템은 현재 진행 중인 이벤트 기간에만 사용할 수 있습니다"
        
        ### 네트워크 오류 처리
        1. 서버 응답 지연
           - 상황: 장착 요청 후 서버 응답이 늦어지는 경우
           - 처리: 로딩 인디케이터 표시, 일정 시간 후 재시도 옵션
           - 메시지: "서버 응답 대기 중... 잠시만 기다려주세요"
        
        2. 동기화 오류
           - 상황: 클라이언트-서버 간 아이템 정보 불일치
           - 처리: 자동 재동기화 시도
           - 메시지: "아이템 정보를 갱신하는 중입니다..."
        
        3. 장착 실패
           - 상황: 서버에서 예기치 않은 이유로 장착 거부
           - 처리: 오류 코드와 함께 실패 메시지 표시
           - 메시지: "아이템 장착에 실패했습니다. (오류 코드: {코드})"
        """
    ]
    
    processed_chunks = []
    chunk_id = 0
    
    for i, text in enumerate(sample_texts):
        # 텍스트 분할
        chunks = split_text(text, 1000, 200)
        
        for j, chunk_text in enumerate(chunks):
            processed_chunks.append({
                'text': chunk_text,
                'metadata': {
                    'file_name': 'sample_game_design_doc.docx',
                    'chunk_id': chunk_id,
                    'section_id': i,
                    'source': 'internal_sample',
                    'type': 'game_design_doc'
                }
            })
            chunk_id += 1
    
    return processed_chunks