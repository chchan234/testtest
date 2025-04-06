"""
문서 처리 모듈: DOCX, PDF 파일을 텍스트로 추출하고 청크 단위로 분리
"""

import os
from typing import List, Dict, Any, Union
import docx
import fitz  # PyMuPDF
from langchain.text_splitter import RecursiveCharacterTextSplitter

def extract_text(file_path: str) -> str:
    """
    DOCX 또는 PDF 파일에서 텍스트를 추출
    
    Args:
        file_path: 처리할 파일 경로
        
    Returns:
        추출된 전체 텍스트
    """
    file_ext = os.path.splitext(file_path)[1].lower()
    
    if file_ext == '.docx':
        return _extract_from_docx(file_path)
    elif file_ext == '.pdf':
        return _extract_from_pdf(file_path)
    else:
        raise ValueError(f"지원하지 않는 파일 형식입니다: {file_ext}")

def _extract_from_docx(file_path: str) -> str:
    """DOCX 파일에서 텍스트를 추출"""
    doc = docx.Document(file_path)
    full_text = []
    
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    # 표(tables)에서 텍스트 추출
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
    
    return '\n'.join(full_text)

def _extract_from_pdf(file_path: str) -> str:
    """PDF 파일에서 텍스트를 추출"""
    doc = fitz.open(file_path)
    full_text = []
    
    for page in doc:
        full_text.append(page.get_text())
    
    return '\n'.join(full_text)

def process_document(file_path: str, chunk_size: int = 1000, chunk_overlap: int = 200) -> List[Dict[str, Any]]:
    """
    문서를 처리하여 청크 단위로 분리
    
    Args:
        file_path: 처리할 파일 경로
        chunk_size: 각 청크의 최대 크기
        chunk_overlap: 청크 간 겹치는 문자 수
        
    Returns:
        청크 리스트 (메타데이터 포함)
    """
    # 텍스트 추출
    text = extract_text(file_path)
    
    # 텍스트 분할
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len
    )
    
    chunks = text_splitter.create_documents([text])
    
    # 메타데이터 추가 (파일명, 페이지 번호 등)
    file_name = os.path.basename(file_path)
    processed_chunks = []
    
    for i, chunk in enumerate(chunks):
        processed_chunks.append({
            'text': chunk.page_content,
            'metadata': {
                'file_name': file_name,
                'chunk_id': i,
                'source': file_path
            }
        })
    
    return processed_chunks